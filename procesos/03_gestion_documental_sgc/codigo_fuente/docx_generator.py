import os
import re
import datetime
import docx
from docxtpl import DocxTemplate
from src import config

def insert_paragraph_after(paragraph, text, style=None):
    """Inserta un párrafo inmediatamente después del párrafo dado en el XML de Word."""
    new_p = docx.oxml.shared.OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if style:
        new_para.style = style
    return new_para

def prepare_docx_template(filepath, output_path, doc_type="Procedimiento"):
    """
    Toma un archivo .docx de plantilla básica y le inyecta las etiquetas Jinja2 
    debajo de cada encabezado estándar, guardándolo como una plantilla Jinja2 válida.
    Esto permite usar plantillas simples descargadas desde Google Docs directamente.
    """
    config.logger.info(f"Preprocesando plantilla básica en {filepath} para insertar etiquetas Jinja2...")
    doc = docx.Document(filepath)
    
    # 1. Buscar y reemplazar el título
    for p in doc.paragraphs:
        p_text_norm = p.text.strip().lower()
        if "título:" in p_text_norm or "t&iacute;tulo:" in p_text_norm or "tìtulo:" in p_text_norm:
            p.text = "Título: {{ titulo }}"
            
    # 2. Insertar placeholders debajo de las cabeceras estándar
    paragraphs_list = list(doc.paragraphs)
    for p in paragraphs_list:
        text_clean = p.text.strip().lower()
        
        if text_clean == "objetivo":
            insert_paragraph_after(p, "{{ objetivo }}")
        elif text_clean == "alcance":
            insert_paragraph_after(p, "{{ alcance }}")
        elif text_clean == "responsabilidades":
            insert_paragraph_after(p, "{% for resp in responsabilidades %}\n• {{ resp }}\n{% endfor %}")
        elif text_clean == "normas":
            insert_paragraph_after(p, "{% for norma in normas %}\n• {{ norma }}\n{% endfor %}")
        elif text_clean in ["definiciones de términos", "definiciones de tèrminos"]:
            insert_paragraph_after(p, "{% for def in definiciones %}\n• {{ def.termino }}: {{ def.definicion }}\n{% endfor %}")
        elif text_clean == "documentos de referencia":
            insert_paragraph_after(p, "{% for doc_ref in documentos_referencia %}\n• {{ doc_ref }}\n{% endfor %}")
        elif text_clean in ["descripción del procedimiento", "descripciòn del procedimiento"]:
            insert_paragraph_after(p, "{% for paso in pasos %}\n**Paso {{ paso.numero }} ({{ paso.responsable }}): {{ paso.actividad }}**\n{{ paso.descripcion }}\n\n{% endfor %}")
        elif text_clean == "flujo del procedimiento":
            insert_paragraph_after(p, "Ver diagrama de flujo Mermaid generado al final del documento.")

    # 3. Modificar el Pie de Página (Footer) dinámicamente
    # El footer contiene 'SGC-CMP-PR-01 Versión 1- 2025 Pag.'
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            if "SGC-" in p.text:
                config.logger.info(f"Encontrado código de nomenclatura en pie de página: '{p.text}'")
                # Reemplazar código fijo (ej. SGC-CMP-PR-01) por {{ codigo }}
                new_text = re.sub(r'SGC-[A-Z]+-[A-Z]+-\d+', '{{ codigo }}', p.text)
                # Reemplazar Versión 1 por Versión {{ version }}
                new_text = re.sub(r'Versin \d+|Versión \d+|Versi&oacute;n \d+', 'Versión {{ version }}', new_text)
                # Reemplazar año 2025 o 2026 por {{ anio }}
                new_text = re.sub(r'2025|2026', '{{ anio }}', new_text)
                p.text = new_text
                config.logger.info(f"Pie de página modificado a: '{p.text}'")

    # 4. Modificar la tabla de encabezado (Cajetín)
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                row_texts = [cell.text.strip().lower() for cell in row.cells]
                has_version_label = any("versi" in text for text in row_texts)
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p_text = p.text.strip()
                        p_text_norm = p_text.lower()
                        
                        # A. Reemplazar título del cajetín
                        if "compras en Enterprise SGC" in p_text_norm or "nombre del documento" in p_text_norm:
                            if p.runs:
                                p.runs[0].text = '{{ titulo }}'
                                for r in p.runs[1:]:
                                    r.text = ''
                            else:
                                p.text = '{{ titulo }}'
                                
                        # B. Reemplazar código del cajetín
                        elif re.match(r'^SGC-[A-Z0-9]+-[A-Z0-9]+-\d+$', p_text):
                            if p.runs:
                                p.runs[0].text = '{{ codigo }}'
                                for r in p.runs[1:]:
                                    r.text = ''
                            else:
                                p.text = '{{ codigo }}'
                                
                        # C. Reemplazar versión del cajetín
                        elif p_text in ['01', '1.0', '1.0.0'] and has_version_label:
                            if p.runs:
                                p.runs[0].text = '{{ version }}'
                                for r in p.runs[1:]:
                                    r.text = ''
                            else:
                                p.text = '{{ version }}'

    doc.save(output_path)
    config.logger.info(f"Plantilla preprocesada y guardada con éxito en {output_path}")

class DocxGenerator:
    @staticmethod
    def generate_document(doc_type: str, data_dict: dict, output_filename: str) -> str:
        """
        Genera el documento final rellenando las plantillas preprocesadas con la información 
        estructurada por la IA.
        
        :param doc_type: 'Procedimiento' o 'Norma'
        :param data_dict: Diccionario de datos que contiene titulo, objetivo, alcance, pasos, etc.
        :param output_filename: Nombre del archivo de salida (ej: 'SGC-INV-PR-17 Procedimiento de Inventario.docx')
        :return: Ruta absoluta del archivo generado.
        """
        # Determinar rutas
        template_name = "plantilla_procedimiento.docx" if doc_type == "Procedimiento" else "plantilla_norma.docx"
        raw_template_path = os.path.join(config.TEMPLATES_DIR, template_name)
        
        # Plantilla intermedia con tags Jinja2
        tagged_template_name = f"tagged_{template_name}"
        tagged_template_path = os.path.join(config.TEMPLATES_DIR, tagged_template_name)
        
        # 1. Asegurar que la plantilla tenga los tags Jinja2 correctos
        prepare_docx_template(raw_template_path, tagged_template_path, doc_type)
        
        # 2. Cargar con docxtpl y renderizar
        doc_tpl = DocxTemplate(tagged_template_path)
        
        # Inyectar variables globales adicionales del cajetín
        anio_actual = str(datetime.datetime.now().year)
        fecha_actual = datetime.datetime.now().strftime("%d/%m/%Y")
        
        context = {
            **data_dict,
            "version": data_dict.get("version", "1.0"),
            "fecha": data_dict.get("fecha", fecha_actual),
            "anio": anio_actual,
            "autor": data_dict.get("autor", "Líder de Procesos")
        }
        
        config.logger.info(f"Renderizando plantilla de {doc_type} con datos dinámicos...")
        doc_tpl.render(context)
        
        # Guardar en la carpeta de salidas
        final_output_path = os.path.join(config.OUTPUTS_DIR, output_filename)
        doc_tpl.save(final_output_path)
        config.logger.info(f"¡Documento generado con éxito! Guardado en {final_output_path}")
        
        return final_output_path

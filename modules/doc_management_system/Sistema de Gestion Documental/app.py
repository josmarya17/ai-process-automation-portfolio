import os
import datetime
import re
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from src import config
from src.inventory_manager import InventoryManager
from src.ai_engine import AIEngine
from src.docx_generator import DocxGenerator
from src.svg_generator import SVGGenerator
from src.drawio_generator import DrawIOGenerator
from src.drawio_parser import DrawIOParser
from src import google_client

# Configuración de página de Streamlit
st.set_page_config(
    page_title="Enterprise SGC BPM Automator",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estética Premium Customizada mediante Inyección de CSS (Light Theme)
st.markdown("""
<style>
    /* Estilos generales y fuentes */
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Space+Grotesk:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
        color: #1E293B !important;
    }
    
    .stApp {
        background-color: #F8FAFC !important;
    }
    
    /* Contenedor principal con Card premium corporativo */
    .glass-card {
        background: #FFFFFF !important;
        border-radius: 16px;
        padding: 28px;
        border: 1px solid #E2E8F0 !important;
        box-shadow: 0 4px 20px 0 rgba(0, 0, 0, 0.02) !important;
        margin-bottom: 24px;
        transition: all 0.3s ease;
    }
    .glass-card:hover {
        border: 1px solid rgba(248, 80, 0, 0.3) !important;
        box-shadow: 0 10px 25px 0 rgba(248, 80, 0, 0.04) !important;
    }
    
    /* Título principal con gradiente de Enterprise SGC */
    .main-title {
        background: linear-gradient(90deg, #F85000 0%, #FF8A00 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent !important;
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 800;
        font-size: 3.25rem;
        margin-bottom: 0px;
        letter-spacing: -1.5px;
    }
    
    .subtitle {
        color: #475569 !important;
        font-size: 1.15rem;
        margin-top: 8px;
        margin-bottom: 35px;
        font-weight: 400;
        letter-spacing: 0.2px;
    }
    
    /* Botones de acción ultra premium corporativos Enterprise SGC */
    .stButton>button {
        background: linear-gradient(135deg, #F85000 0%, #FF6A00 100%) !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 14px 28px !important;
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 600 !important;
        font-size: 1.05rem !important;
        box-shadow: 0 4px 15px rgba(248, 80, 0, 0.2) !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 20px rgba(248, 80, 0, 0.35) !important;
        background: linear-gradient(135deg, #FF6A00 0%, #FF8A00 100%) !important;
    }
    
    .stButton>button:active {
        transform: translateY(0px) !important;
    }
    
    /* Nomenclatura Badge con brillo naranja */
    .nomenclature-badge {
        background: rgba(248, 80, 0, 0.05) !important;
        border: 1.5px solid #F85000 !important;
        color: #F85000 !important;
        border-radius: 20px;
        padding: 10px 18px;
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 700;
        font-size: 1.35rem;
        text-align: center;
        margin: 14px 0px;
        box-shadow: 0 2px 10px rgba(248, 80, 0, 0.05);
    }
    
    /* Personalización de Inputs, Selectboxes y TextAreas */
    div[data-baseweb="select"] > div {
        background-color: #FFFFFF !important;
        border: 1px solid #CBD5E1 !important;
        border-radius: 8px !important;
        color: #0F172A !important;
    }
    
    div[data-baseweb="textarea"] > div {
        background-color: #FFFFFF !important;
        border: 1px solid #CBD5E1 !important;
        border-radius: 10px !important;
        color: #0F172A !important;
    }
    
    div[data-baseweb="input"] > div {
        background-color: #FFFFFF !important;
        border: 1px solid #CBD5E1 !important;
        border-radius: 8px !important;
        color: #0F172A !important;
    }
    
    /* Personalización de Pestañas (Tabs) */
    button[data-baseweb="tab"] {
        color: #475569 !important;
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 500 !important;
        background-color: transparent !important;
        border: none !important;
        transition: all 0.3s ease !important;
    }
    
    button[data-baseweb="tab"][aria-selected="true"] {
        color: #F85000 !important;
        font-weight: 700 !important;
        border-bottom: 2px solid #F85000 !important;
    }
    
    /* Estilos del Sidebar Claro */
    [data-testid="stSidebar"] {
        background-color: #F1F5F9 !important;
        border-right: 1px solid #E2E8F0 !important;
    }
    [data-testid="stSidebar"] * {
        color: #1E293B !important;
    }
    [data-testid="stSidebar"] h3 {
        color: #0F172A !important;
    }
    
    /* Override global text for light theme stMarkdown */
    .stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown span {
        color: #1E293B !important;
    }
    
    /* Headers inside main panel */
    h1, h2, h3, h4, h5, h6 {
        color: #0F172A !important;
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 700 !important;
    }
</style>
""", unsafe_allow_html=True)

# Inicializar managers en estado de sesión para persistencia
if 'im' not in st.session_state:
    with st.spinner("Cargando inventario de nomenclatura actual..."):
        st.session_state.im = InventoryManager()

if 'ai' not in st.session_state:
    st.session_state.ai = AIEngine()

im = st.session_state.im
ai = st.session_state.ai

# --- SIDEBAR: ESTADO Y CONFIGURACIÓN ---
with st.sidebar:
    # Mostrar el logotipo oficial de Enterprise SGC si existe
    if os.path.exists("Logo_Enterprise SGC.jpg"):
        st.image("Logo_Enterprise SGC.jpg", use_container_width=True)
    else:
        st.markdown("## ⚡ Enterprise SGC BPM Automator")
    st.markdown("---")
    
    # Validar si existe token.json
    has_oauth = os.path.exists(config.TOKEN_FILE)
    if not has_oauth:
        st.warning("🟡 Modo Fallback (Caché/Offline)")
        st.markdown(
            "<small>El sistema opera usando los CSV públicos de lectura y guardará "
            "los nuevos registros localmente hasta conectar OAuth.</small>",
            unsafe_allow_html=True
        )

    st.markdown("### 🛠️ Parámetros del Documento")
    
    # Selectores dinámicos alimentados por el InventoryManager
    areas = im.get_areas()
    # Limpiar campos no deseados de la hoja de nomenclatura
    areas = [a for a in areas if a not in ['Empresa', 'Sistema']]
    
    selected_area = st.selectbox(
        "Selecciona el Área del Proceso:",
        options=areas,
        index=areas.index("Inventario") if "Inventario" in areas else 0
    )
    
    selected_type = st.selectbox(
        "Tipo de Documento:",
        options=im.get_document_types(),
        index=0
    )
    
    # Calcular y mostrar código sugerido en tiempo real
    suggested_code = im.suggest_next_code(selected_area, selected_type)
    
    # Manejar estado para permitir edición manual de la nomenclatura
    if 'last_selected_area' not in st.session_state or st.session_state.last_selected_area != selected_area or \
       'last_selected_type' not in st.session_state or st.session_state.last_selected_type != selected_type:
        st.session_state.last_selected_area = selected_area
        st.session_state.last_selected_type = selected_type
        st.session_state.document_code = suggested_code
        
    st.markdown("### 🏷️ Nomenclatura Sugerida (Editable)")
    
    # Entrada de texto para que el código sea editable
    edited_code = st.text_input(
        "Código del Documento:",
        value=st.session_state.document_code,
        label_visibility="collapsed",
        help="Código sugerido automáticamente según el área y tipo. Puedes editarlo libremente si es necesario."
    )
    
    # Guardar en estado de sesión el código actual (ya sea sugerido o editado por el usuario)
    st.session_state.document_code = edited_code
    
    # Si la nomenclatura fue editada, mostrar un pequeño indicador
    if edited_code != suggested_code:
        st.markdown(f'<div style="color: #FF8A00; font-size: 0.85rem; margin-top: -10px; margin-bottom: 10px; font-weight: 500;">⚠️ Modificado manualmente (Sugerido: {suggested_code})</div>', unsafe_allow_html=True)
    
    # Nombre / Título del Proceso o Norma (Editable)
    st.markdown("### 📋 Nombre del Proceso / Norma")
    if 'current_doc' in st.session_state:
        default_title = st.session_state.current_doc.titulo
    else:
        default_title = ""
    edited_title = st.text_input(
        "Nombre del Proceso / Norma:",
        value=default_title,
        label_visibility="collapsed",
        help="Título formal que tendrá el documento. Se autocompleta al generar, pero puedes editarlo."
    )
    st.session_state.document_title = edited_title

    # Código Relacionado
    related_code = ""
    if selected_type == "Procedimiento":
        suggested_related = edited_code.replace("-PR-", "-FT-").replace("-IT-", "-FT-")
        st.markdown("### 🔗 Código de Flujo Relacionado")
        related_code = st.text_input(
            "Código de Flujo Relacionado:",
            value=suggested_related,
            label_visibility="collapsed",
            help="Código del flujo de trabajo asociado a este procedimiento."
        )
    elif selected_type == "Flujo de Trabajo":
        suggested_related = edited_code.replace("-FT-", "-PR-")
        st.markdown("### 🔗 Código de Proceso Relacionado")
        related_code = st.text_input(
            "Código de Proceso Relacionado:",
            value=suggested_related,
            label_visibility="collapsed",
            help="Código del procedimiento asociado a este flujo."
        )
    st.session_state.related_code = related_code

    st.markdown("### 📝 Datos de Emisión")
    input_autor = st.text_input("Autor / Redactor:", value="Josmary Pinto")
    input_version = st.text_input("Versión Inicial:", value="1.0")
    fecha_proceso = st.date_input("Fecha de Aprobación:", value=datetime.date.today())
    
    st.markdown("---")
    st.markdown("<small>Hecho por la Líder de Procesos Josmary Pinto & Antigravity</small>", unsafe_allow_html=True)

# --- PANEL CENTRAL: INTERFAZ PRINCIPAL ---
st.markdown('<h1 class="main-title">Automatizador de Procesos BPM</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Transformación inteligente de notas estructuradas en estándares de calidad corporativos</p>', unsafe_allow_html=True)

# Pestañas principales de la aplicación a nivel superior
tab_main_generator, tab_matrix_inventory = st.tabs([
    "⚡ Generador BPM & Importación",
    "🗂️ Matriz de Documentos (Inventario)"
])

# Inicializar variables para procesar
btn_analyze = False
btn_generate_drawio = False
parsed_data = None
raw_input = ""

with tab_main_generator:
    # Inicializar chat si no existe
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = [
            {"role": "assistant", "content": "¡Hola! Soy Enterprise SGC Procesos IA, tu consultora de procesos BPM para Farmacia Enterprise SGC. ⚡\n\n¿Qué proceso corporativo te gustaría relevar, definir o documentar hoy? Si tienes apuntes, minutas o un manual en PDF, puedes subirlo al chat para que lo analicemos e iteremos juntos."}
        ]
        
    if 'processed_pdfs' not in st.session_state:
        st.session_state.processed_pdfs = set()
    if 'processed_files' not in st.session_state:
        st.session_state.processed_files = st.session_state.processed_pdfs

    col_work, col_results = st.columns([5, 7])
    
    with col_work:
        tab_chat, tab_drawio = st.tabs([
            "💬 Enterprise SGC Procesos IA",
            "📥 Importar Draw.io (.drawio)"
        ])
        
        with tab_chat:
            st.markdown('<div class="glass-card" style="margin-bottom:15px; border:none; padding:10px;">', unsafe_allow_html=True)
            st.markdown("### 💬 Conversación con Enterprise SGC Procesos IA")
            st.markdown("<p style='color:#8f9cae; font-size:0.9rem;'>Diseña tu proceso conversando, cargando minutas o subiendo PDFs en el chat:</p>", unsafe_allow_html=True)
            
            # Caja de chat con scroll
            chat_container = st.container(height=380)
            with chat_container:
                for msg in st.session_state.chat_history:
                    display_content = msg["content"]
                    if ("📄 **[Archivo PDF Cargado]**" in display_content or "📄 **[Archivo Word (DOCX) Cargado]**" in display_content) and len(display_content) > 1000:
                        lines = display_content.split("\n")
                        file_header = lines[0]
                        display_content = f"{file_header}\n\n*(Texto completo del archivo cargado en el contexto del chat para análisis)*"
                        
                    with st.chat_message(msg["role"]):
                        st.markdown(display_content)
            
            # Cargador de Archivos (PDF y Word) - Múltiple
            uploaded_files = st.file_uploader(
                "📄 Adjuntar archivos al chat (PDF, Word / .docx):",
                type=["pdf", "docx"],
                accept_multiple_files=True,
                key="chat_files_uploader"
            )
            
            pending_files = []
            if uploaded_files:
                pending_files = [f for f in uploaded_files if f.name not in st.session_state.processed_files]
                if pending_files:
                    st.markdown(
                        f'<div style="background-color: rgba(248, 80, 0, 0.05); border: 1.2px dashed #F85000; border-radius: 8px; padding: 12px; margin-bottom: 15px;">'
                        f'<span style="font-weight: 700; color: #F85000;">📎 Archivos seleccionados (Listos para enviar con tu mensaje):</span><br>'
                        + "<br>".join([f'<span style="font-size: 0.9rem; color: #1E293B;">• 📄 {f.name}</span>' for f in pending_files])
                        + '</div>',
                        unsafe_allow_html=True
                    )
                    
                    # Botón opcional para enviar los archivos solos de inmediato
                    col_send_files, _ = st.columns([2, 1])
                    with col_send_files:
                        if st.button("📤 Enviar archivos adjuntos solos", key="btn_send_files_alone", use_container_width=True):
                            combined_message_parts = []
                            success_count = 0
                            
                            with st.spinner("Procesando y extrayendo contenido de los archivos cargados..."):
                                for f in pending_files:
                                    try:
                                        file_text = ""
                                        if f.name.lower().endswith(".pdf"):
                                            import pypdf
                                            reader = pypdf.PdfReader(f)
                                            pdf_text = ""
                                            for page in reader.pages:
                                                text = page.extract_text()
                                                if text:
                                                    pdf_text += text + "\n"
                                            file_text = pdf_text.strip()
                                            doc_type_label = "PDF"
                                        elif f.name.lower().endswith(".docx"):
                                            import docx
                                            doc_obj = docx.Document(f)
                                            text_parts = []
                                            for p in doc_obj.paragraphs:
                                                if p.text.strip():
                                                    text_parts.append(p.text.strip())
                                            for table in doc_obj.tables:
                                                for row in table.rows:
                                                    row_text = []
                                                    for cell in row.cells:
                                                        cell_text = cell.text.strip()
                                                        if cell_text and cell_text not in row_text:
                                                            row_text.append(cell_text)
                                                    if row_text:
                                                        text_parts.append(" | ".join(row_text))
                                            file_text = "\n".join(text_parts).strip()
                                            doc_type_label = "Word (DOCX)"
                                        
                                        if file_text:
                                            file_msg = f"📄 **[Archivo {doc_type_label} Cargado]** `{f.name}`\n\nContenido:\n{file_text}"
                                            st.session_state.chat_history.append({"role": "user", "content": file_msg})
                                            combined_message_parts.append(f"He cargado el archivo {doc_type_label} '{f.name}'.")
                                            st.session_state.processed_files.add(f.name)
                                            success_count += 1
                                        else:
                                            st.error(f"No se pudo extraer texto del archivo {f.name}.")
                                    except Exception as e:
                                        st.error(f"Error al procesar el archivo {f.name}: {e}")
                            
                            if success_count > 0:
                                with st.spinner("El Consultor BPM está analizando la información de los archivos..."):
                                    summary_trigger = "\n".join(combined_message_parts) + "\nPor favor, analiza la información de estos nuevos archivos para definir el procedimiento."
                                    try:
                                        response = ai.get_chat_response(
                                            chat_history=st.session_state.chat_history, # Incluir los archivos en el historial
                                            user_message=summary_trigger,
                                            doc_type=selected_type
                                        )
                                        st.session_state.chat_history.append({"role": "assistant", "content": response})
                                    except Exception as e:
                                        st.error(f"Error al obtener respuesta de Gemini: {e}")
                                st.rerun()
            
            # Entrada de texto del chat
            user_input = st.chat_input("Escribe tus instrucciones o notas aquí...")
            if user_input:
                success_count = 0
                combined_message_parts = []
                
                # Procesar archivos pendientes antes si los hay
                if pending_files:
                    with st.spinner("Procesando y extrayendo contenido de los archivos cargados..."):
                        for f in pending_files:
                            try:
                                file_text = ""
                                if f.name.lower().endswith(".pdf"):
                                    import pypdf
                                    reader = pypdf.PdfReader(f)
                                    pdf_text = ""
                                    for page in reader.pages:
                                        text = page.extract_text()
                                        if text:
                                            pdf_text += text + "\n"
                                    file_text = pdf_text.strip()
                                    doc_type_label = "PDF"
                                elif f.name.lower().endswith(".docx"):
                                    import docx
                                    doc_obj = docx.Document(f)
                                    text_parts = []
                                    for p in doc_obj.paragraphs:
                                        if p.text.strip():
                                            text_parts.append(p.text.strip())
                                    for table in doc_obj.tables:
                                        for row in table.rows:
                                            row_text = []
                                            for cell in row.cells:
                                                cell_text = cell.text.strip()
                                                if cell_text and cell_text not in row_text:
                                                    row_text.append(cell_text)
                                            if row_text:
                                                text_parts.append(" | ".join(row_text))
                                    file_text = "\n".join(text_parts).strip()
                                    doc_type_label = "Word (DOCX)"
                                
                                if file_text:
                                    file_msg = f"📄 **[Archivo {doc_type_label} Cargado]** `{f.name}`\n\nContenido:\n{file_text}"
                                    st.session_state.chat_history.append({"role": "user", "content": file_msg})
                                    combined_message_parts.append(f"He cargado el archivo {doc_type_label} '{f.name}'.")
                                    st.session_state.processed_files.add(f.name)
                                    success_count += 1
                                else:
                                    st.error(f"No se pudo extraer texto del archivo {f.name}.")
                            except Exception as e:
                                st.error(f"Error al procesar el archivo {f.name}: {e}")
                
                # Agregar el mensaje del usuario
                st.session_state.chat_history.append({"role": "user", "content": user_input})
                
                # Obtener la respuesta de Gemini en base al historial acumulado
                with st.spinner("El Consultor BPM está respondiendo..."):
                    try:
                        # Pasamos todo el historial excluyendo el último mensaje de texto del usuario (el cual es user_message)
                        response = ai.get_chat_response(
                            chat_history=st.session_state.chat_history[:-1],
                            user_message=user_input,
                            doc_type=selected_type
                        )
                        st.session_state.chat_history.append({"role": "assistant", "content": response})
                    except Exception as e:
                        st.error(f"Error al obtener respuesta de Gemini: {e}")
                st.rerun()
            
            # Ajustar la etiqueta y mensaje de consolidación según el tipo seleccionado
            if selected_type == "Flujo de Trabajo":
                chat_btn_label = "⚡ Generar Flujograma BPM"
                spinner_msg = "Consolidando toda la conversación en un diagrama de carriles estilo Figma..."
                success_msg = "¡Flujograma de proceso consolidado y diagramas vectoriales generados con éxito!"
            elif selected_type == "Norma":
                chat_btn_label = "⚡ Generar Estructura de Norma"
                spinner_msg = "Consolidando toda la conversación en un documento de Norma..."
                success_msg = "¡Estructura de la Norma consolidada con éxito!"
            elif selected_type == "Procedimiento":
                chat_btn_label = "⚡ Generar Procedimiento"
                spinner_msg = "Consolidando toda la conversación en un diagrama de carriles estilo Figma y documento Word..."
                success_msg = "¡Procedimiento y flujograma consolidados con éxito!"
            elif selected_type == "Instrucción de Trabajo":
                chat_btn_label = "⚡ Generar Instrucción de Trabajo"
                spinner_msg = "Consolidando toda la conversación en un instructivo de trabajo..."
                success_msg = "¡Instrucción de trabajo consolidada con éxito!"
            else:
                chat_btn_label = "⚡ Generar Estructura y Flujograma BPM"
                spinner_msg = "Consolidando toda la conversación en un diagrama de carriles estilo Figma y documento Word..."
                success_msg = "¡Estructura de proceso consolidada y diagramas vectoriales generados con éxito!"

            # Botones de Acción de Chat
            col_chat_btns = st.columns([2, 1])
            with col_chat_btns[0]:
                btn_generate_from_chat = st.button(chat_btn_label, key="btn_gen_chat", use_container_width=True)
            with col_chat_btns[1]:
                btn_clear_chat = st.button("🗑️ Reiniciar Chat", key="btn_clear_chat", use_container_width=True)
                
            if btn_clear_chat:
                st.session_state.chat_history = [
                    {"role": "assistant", "content": "¡Hola! Soy Enterprise SGC Procesos IA, tu consultora de procesos BPM para Farmacia Enterprise SGC. ⚡\n\n¿Qué proceso corporativo te gustaría relevar, definir o documentar hoy? Si tienes apuntes, minutas o un manual en PDF o documento Word, puedes subirlo al chat para que lo analicemos e iteremos juntos."}
                ]
                st.session_state.processed_pdfs = set()
                st.session_state.processed_files = st.session_state.processed_pdfs
                if 'current_doc' in st.session_state:
                    del st.session_state.current_doc
                st.rerun()
                
            if btn_generate_from_chat:
                with st.spinner(spinner_msg):
                    try:
                        structured_doc = ai.analyze_chat_conversation(st.session_state.chat_history, selected_type)
                        st.session_state.current_doc = structured_doc
                        st.session_state.nomenclatura_asignada = st.session_state.get('document_code', suggested_code)
                        st.success(success_msg)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al consolidar la conversación: {e}")
            st.markdown('</div>', unsafe_allow_html=True)
            
        with tab_drawio:
            st.markdown('<div class="glass-card" style="margin-bottom:15px; border:none; padding:10px;">', unsafe_allow_html=True)
            st.markdown("### 📥 Importar Diagrama desde Draw.io / Diagrams.net")
            st.markdown("<p style='color:#8f9cae; font-size:0.9rem;'>Sube tu archivo .drawio editado o exportado para reconstruir su procedimiento de forma automática:</p>", unsafe_allow_html=True)
            
            uploaded_drawio = st.file_uploader(
                "Selecciona tu archivo de diagrama (.drawio o .xml):",
                type=["drawio", "xml"],
                key="drawio_file_uploader",
                label_visibility="collapsed"
            )
            
            if uploaded_drawio is not None:
                try:
                    drawio_content = uploaded_drawio.read().decode("utf-8")
                    parsed_data = DrawIOParser.parse_drawio(drawio_content)
                    
                    if "error" in parsed_data:
                        st.error(parsed_data["error"])
                    else:
                        st.success("¡Archivo de Draw.io cargado y procesado con éxito!")
                        
                        with st.expander("🔍 Vista Previa del Diagrama Extraído"):
                            st.markdown("**Pasos / Actividades Detectadas:**")
                            for step in parsed_data.get("steps", []):
                                st.markdown(f"• **{step['role']}**: {step['text']}")
                                
                            if parsed_data.get("connections"):
                                st.markdown("---")
                                st.markdown("**Conexiones / Flujo Lógico:**")
                                for conn in parsed_data.get("connections", []):
                                    st.markdown(f"• {conn}")
                        
                        btn_generate_drawio = st.button("⚡ Generar Documento desde Draw.io", key="btn_generate_draw", use_container_width=True)
                        if btn_generate_drawio:
                            with st.spinner("La IA de Gemini está analizando tu diagrama de Draw.io y generando el procedimiento..."):
                                try:
                                    structured_doc = ai.generate_from_drawio(parsed_data, selected_type)
                                    st.session_state.current_doc = structured_doc
                                    st.session_state.nomenclatura_asignada = st.session_state.get('document_code', suggested_code)
                                    st.success("¡Procedimiento y flujograma reconstruidos con éxito desde Draw.io!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Error al generar desde el diagrama de Draw.io: {e}")
                except Exception as e:
                    st.error(f"Error al leer o parsear el archivo .drawio: {e}")
            st.markdown('</div>', unsafe_allow_html=True)

    with col_results:
        if 'current_doc' in st.session_state:
            doc = st.session_state.current_doc
            code = st.session_state.nomenclatura_asignada
            title_to_display = st.session_state.get('document_title', doc.titulo)
            if not title_to_display.strip():
                title_to_display = doc.titulo
            
            st.markdown(f"## 📋 Proceso: {title_to_display} ({code})")
            
            # Tabs de visualización de resultados según tipo de documento
            if selected_type == "Flujo de Trabajo":
                tab_mermaid, tab_steps = st.tabs([
                    "🟠 Flujograma de Carriles",
                    "📋 Matriz de Actividades"
                ])
            elif selected_type == "Norma":
                tab_doc = st.tabs([
                    "📄 Estructura de Calidad"
                ])[0]
            else:
                tab_doc, tab_mermaid, tab_steps = st.tabs([
                    "📄 Estructura de Calidad",
                    "🟠 Flujograma de Carriles",
                    "📋 Matriz de Actividades"
                ])
            
            # TAB 1: Documento Estructurado
            if selected_type != "Flujo de Trabajo":
                with tab_doc:
                    col_doc_left, col_doc_right = st.columns([2, 1])
                    
                    with col_doc_left:
                        st.markdown(f"### 🎯 Objetivo\n{doc.objetivo}")
                        st.markdown(f"### 📍 Alcance\n{doc.alcance}")
                        
                        st.markdown("### 📝 Normas y Políticas del Proceso")
                        for n in doc.normas:
                            st.markdown(f"• {n}")
                            
                        st.markdown("### 🔗 Documentos y Sistemas de Referencia")
                        for r in doc.documentos_reference if hasattr(doc, 'documentos_reference') else doc.documentos_referencia:
                            st.markdown(f"• {r}")
                            
                    with col_doc_right:
                        st.markdown("### 👥 Responsabilidades Macro")
                        for r in doc.responsabilidades:
                            st.markdown(f"**{r.split(':')[0] if ':' in r else 'Rol'}:** {r.split(':')[1] if ':' in r else r}")
                            
                        st.markdown("### 📖 Glosario de Definiciones")
                        for d in doc.definiciones:
                            st.markdown(f"• **{d.termino}**: {d.definicion}")

            # TAB 2: Diagrama Vectorial (SVG)
            if selected_type in ["Procedimiento", "Instrucción de Trabajo", "Flujo de Trabajo"]:
                with tab_mermaid:
                    st.markdown("### 🟠 Flujograma de Carriles (Metodología Enterprise SGC)")
                    st.markdown("<p style='color:#64748B; font-size:0.95rem; margin-top:-10px; margin-bottom: 20px;'>Lienzo vectorial nativo en alta resolución, alineado con el cajetín oficial y listo para usar en tus reportes o editar en Figma:</p>", unsafe_allow_html=True)
                    
                    # Generar el archivo SVG vectorial nativo y el editable de Draw.io
                    fecha_str_h = fecha_proceso.strftime("%d/%m/%Y")
                    svg_filepath = SVGGenerator.generate_svg(doc, code, input_version, fecha_str_h)
                    drawio_filepath = DrawIOGenerator.generate_drawio(doc, code, input_version, fecha_str_h)
                    
                    # Mostrar el SVG directamente de forma fluida, responsive y sin iframes/scrollbars molestos
                    st.markdown('<div class="glass-card" style="padding: 10px; background-color: #FAFAFC !important; border: 1.5px solid #E2E8F0 !important;">', unsafe_allow_html=True)
                    st.image(svg_filepath, use_container_width=True)
                    st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Botones de exportación y descarga
                    col_dl_1, col_dl_2 = st.columns([1, 1])
                    with col_dl_1:
                        with open(svg_filepath, "rb") as file:
                            st.download_button(
                                label="⬇️ Descargar SVG (Figma)",
                                data=file,
                                file_name=f"{code}_Diagrama_Enterprise SGC.svg",
                                mime="image/svg+xml",
                                use_container_width=True
                            )
                    with col_dl_2:
                        with open(drawio_filepath, "rb") as file:
                            st.download_button(
                                label="⚡ Descargar Editable para Draw.io (.drawio)",
                                data=file,
                                file_name=f"{code}_Diagrama_Enterprise SGC.drawio",
                                mime="application/octet-stream",
                                use_container_width=True
                            )
                    
                    # Manual de Integración Premium
                    st.markdown("---")
                    with st.expander("⚡ 💻 CÓMO ABRIR Y EDITAR FÁCILMENTE EN DRAW.IO (Recomendado - 100% Ajustable)"):
                        st.markdown("""
                        ### 🚀 ¡Tu diagrama ahora es totalmente interactivo y auto-ajustable!
                        Si encuentras que editar en Figma es molesto, **Draw.io (diagrams.net)** es la solución perfecta, profesional y ágil. Es **100% gratuito**, no requiere registro y funciona de forma local u online.
                        
                        Al usar el archivo **`.drawio`** descargado:
                        1. Entra a **[draw.io](https://app.diagrams.net/)** (o abre Draw.io Desktop si lo tienes instalado).
                        2. Selecciona **"Abrir diagrama existente" (Open Existing Diagram)** y sube el archivo `.drawio` que acabas de descargar.
                        3. **¡Y a diseñar de forma inteligente!**
                           * 🔗 **Flechas Conectadas e Inteligentes**: Si mueves o arrastras cualquier caja a otra posición, las flechas **se doblan y se mantienen conectadas automáticamente** (estilo Visio). ¡Nunca más tendrás que rediseñar flechas o alineaciones rotas!
                           * ✍️ **Textos 100% Editables**: Solo haz **doble clic** sobre el texto de cualquier caja de actividad o carril para reescribirlo en segundos.
                           * 🎨 **Estructura Enterprise SGC Completa**: Se conservan los colores oficiales de Farmacia Enterprise SGC (Naranja, Melocotón, Gris, Blanco), la división inferior de "Sistema / Manual", los carriles punteados y el **Logotipo Oficial de Enterprise SGC** con fidelidad total.
                        """)

                    with st.expander("🎨 Cómo editar en Figma / FigJam (Vectorial)"):
                        st.markdown("""
                        ### 🚀 Importación en Figma
                        Los archivos **`.jam`** son un formato binario y propietario cerrado de Figma. Ninguna aplicación externa puede generarlos directamente sin romper la compatibilidad de la herramienta. Por ello, el estándar profesional de Figma/FigJam es importar a través de **SVG Vectorial Nativo**.
                        
                        Para obtener el diagrama vectorial **100% editable**, sigue estos sencillos pasos:
                        
                        1. Haz clic en el botón de arriba **"⬇️ Descargar SVG (Figma)"** y guarda el archivo `.svg` en tu computadora.
                        2. Abre tu lienzo de **Figma** o **FigJam** (puedes arrastrarlo directamente al archivo `plantilla ejemplo.jam` que tienes en tu carpeta).
                        3. **Arrastra el archivo `.svg` descargado y suéltalo** en cualquier parte del lienzo de Figma.
                        4. ⚠️ **PASO CLAVE PARA EDITAR**: Al importarse, Figma agrupará el diagrama como un único bloque. Para editar libremente cada elemento, simplemente haz **Clic derecho en el bloque importado y selecciona "Ungroup" (Desagrupar)** o presiona el atajo de teclado:
                           * **Windows:** `Ctrl + Shift + G`
                           * **Mac:** `Cmd + Shift + G`
                        5. **¡Y listo!** Cada caja de actividad, texto, conector o etiqueta ahora es un elemento nativo editable de Figma. 
                           * Los textos de las cajas estarán limpios, sin códigos HTML y con sus saltos de línea perfectos.
                           * El **Logotipo Oficial de Enterprise SGC** y los colores corporativos se conservan en alta resolución vectorizada.
                           * Podrás mover, reordenar, cambiar de color o editar cualquier caja o conector libremente.
                        """)
                        
                    with st.expander("🛠️ Ver Código Mermaid.js (Compatible con Visio y draw.io)"):
                        st.markdown("""
                        Este código utiliza saltos de línea reales (`\\n`) y es **100% libre de códigos HTML crudos** para asegurar compatibilidad directa al importar diagramas en herramientas como Microsoft Visio, draw.io o Mermaid Live Editor:
                        """)
                        st.code(doc.mermaid_code, language="mermaid")
                    
                    # LEYENDA BPMN FIGMA
                    st.markdown("---")
                    st.markdown("### 🗺️ Leyenda de Notación Oficial (Estilo Figma)")
                    
                    legend_html = """
                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 16px; margin-top: 15px; font-family: 'Outfit', sans-serif;">
                        <!-- INICIO / FIN -->
                        <div style="background: rgba(22, 26, 38, 0.4); border: 1px solid rgba(255,255,255,0.05); padding: 14px; border-radius: 10px; display: flex; align-items: center; gap: 14px;">
                            <div style="background-color: #E9ECEF; border: 1.5px solid #ADB5BD; border-radius: 20px; min-width: 70px; height: 35px; display: flex; align-items: center; justify-content: center; font-size: 0.75rem; font-weight: bold; color: #000;">Inicio</div>
                            <div>
                                <div style="font-weight: 600; font-size: 0.95rem; color: #FF8A00;">Inicio / Fin</div>
                                <div style="color: #94A3B8; font-size: 0.75rem; margin-top: 2px;">Óvalos grises que demarcan el inicio y el término del flujo de procesos.</div>
                            </div>
                        </div>
                        <!-- AUTOMATICA / SISTEMA -->
                        <div style="background: rgba(22, 26, 38, 0.4); border: 1px solid rgba(255,255,255,0.05); padding: 14px; border-radius: 10px; display: flex; align-items: center; gap: 14px;">
                            <div style="background-color: #FFEADB; border: 1.5px solid #000000; border-radius: 6px; min-width: 80px; height: 45px; display: flex; flex-direction: column; align-items: center; justify-content: center; font-size: 0.7rem; color: #000; font-weight: bold; padding: 2px; box-sizing: border-box;">
                                <div style="font-size: 0.65rem; line-height: 1;">Actividad</div>
                                <hr style="border:0; border-top:1.5px solid #000; width: 100%; margin: 4px 0 2px 0;">
                                <div style="font-size: 0.55rem; color: #F85000; line-height: 1; font-weight: 900;">Sistema</div>
                            </div>
                            <div>
                                <div style="font-weight: 600; font-size: 0.95rem; color: #FF8A00;">Actividad en Sistema</div>
                                <div style="color: #94A3B8; font-size: 0.75rem; margin-top: 2px;">Caja con fondo salmón y división. Proceso automatizado en Odoo u otro software.</div>
                            </div>
                        </div>
                        <!-- MANUAL -->
                        <div style="background: rgba(22, 26, 38, 0.4); border: 1px solid rgba(255,255,255,0.05); padding: 14px; border-radius: 10px; display: flex; align-items: center; gap: 14px;">
                            <div style="background-color: #FFFFFF; border: 1.5px solid #000000; border-radius: 6px; min-width: 80px; height: 45px; display: flex; flex-direction: column; align-items: center; justify-content: center; font-size: 0.7rem; color: #000; font-weight: bold; padding: 2px; box-sizing: border-box;">
                                <div style="font-size: 0.65rem; line-height: 1;">Actividad</div>
                                <hr style="border:0; border-top:1.5px solid #000; width: 100%; margin: 4px 0 2px 0;">
                                <div style="font-size: 0.55rem; color: #555; line-height: 1;">Manual</div>
                            </div>
                            <div>
                                <div style="font-weight: 600; font-size: 0.95rem; color: #FF8A00;">Actividad Manual</div>
                                <div style="color: #94A3B8; font-size: 0.75rem; margin-top: 2px;">Caja blanca con división. Tareas manuales, firmas físicas o correos directos.</div>
                            </div>
                        </div>
                        <!-- DECISION -->
                        <div style="background: rgba(22, 26, 38, 0.4); border: 1px solid rgba(255,255,255,0.05); padding: 14px; border-radius: 10px; display: flex; align-items: center; gap: 14px;">
                            <div style="min-width: 50px; display: flex; align-items: center; justify-content: center;">
                                <div style="background-color: #FFFFFF; border: 1.5px solid #000000; width: 32px; height: 32px; transform: rotate(45deg); display: flex; align-items: center; justify-content: center; box-sizing: border-box;">
                                    <div style="transform: rotate(-45deg); font-size: 0.55rem; font-weight: bold; color: #000;">¿?</div>
                                </div>
                            </div>
                            <div>
                                <div style="font-weight: 600; font-size: 0.95rem; color: #FF8A00;">Pregunta / Decisión</div>
                                <div style="color: #94A3B8; font-size: 0.75rem; margin-top: 2px;">Rombos de decisión. Evalúan condiciones lógicas con caminos de 'Sí' o 'No'.</div>
                            </div>
                        </div>
                        <!-- DOCUMENTO / ANEXO -->
                        <div style="background: rgba(22, 26, 38, 0.4); border: 1px solid rgba(255, 255, 255, 0.05); padding: 14px; border-radius: 10px; display: flex; align-items: center; gap: 14px;">
                            <div style="background-color: #F8FAFC; border: 1.5px dashed #94A3B8; border-radius: 6px; min-width: 80px; height: 35px; display: flex; align-items: center; justify-content: center; font-size: 0.65rem; color: #475569; font-weight: bold;">📄 Anexo</div>
                            <div>
                                <div style="font-weight: 600; font-size: 0.95rem; color: #FF8A00;">Anexo / Documento</div>
                                <div style="color: #94A3B8; font-size: 0.75rem; margin-top: 2px;">Formatos, facturas, plantillas y archivos con bordes punteados.</div>
                            </div>
                        </div>
                    </div>
                    """
                    st.markdown(legend_html, unsafe_allow_html=True)
                    
                    with st.expander("🎨 Importar a Figma / FigJam (Vectorial y Editable)"):
                        st.markdown("""
                        ### 🚀 ¿Cómo importar este flujograma a Figma o FigJam?
                        FigJam y Figma permiten convertir este código Mermaid en diagramas vectoriales **100% editables** en cuestión de segundos:
                        
                        1. **Copia el código Mermaid** que está en la caja de texto inferior (puedes usar el botón de copiar incorporado).
                        2. Abre tu archivo en **FigJam** (como la plantilla `plantilla ejemplo.jam` que has pegado en la carpeta, o un tablero nuevo).
                        3. En la barra de herramientas inferior, haz clic en el icono de **Recursos (Resources)** (o presiona `Shift + I`).
                        4. Ve a la pestaña **Plugins** y busca **"Mermaid"** (por ejemplo, el plugin oficial *"Mermaid"* o *"Mermaid to Diagram"*).
                        5. Ejecuta el plugin, **pega el código** en el cuadro de texto y haz clic en **"Generate"** o **"Create"**.
                        6. **¡Listo!** FigJam creará al instante todo el flujograma vectorial con sus respectivos carriles de roles, cajas de actividades manuales/sistema diferenciadas y la paleta de colores premium de **Enterprise SGC** (Naranja `#F85000`, Salmón `#FFEADB`, Gris `#E9ECEF` y Blanco). ¡Totalmente editable y escalable!
                        """)
                        st.code(doc.mermaid_code, language="mermaid")

            # TAB 3: Pasos detallados
            if selected_type in ["Procedimiento", "Instrucción de Trabajo", "Flujo de Trabajo"]:
                with tab_steps:
                    st.markdown("### 📋 Secuencia de Actividades")
                    
                    # Convertir lista de pasos en Pandas DataFrame
                    steps_data = []
                    for p in doc.pasos:
                        steps_data.append({
                            "Nº": p.numero,
                            "Actividad": p.actividad,
                            "Responsable": p.responsable,
                            "Descripción Detallada": p.descripcion
                        })
                    df_steps = pd.DataFrame(steps_data)
                    
                    # Mostrar tabla premium
                    st.dataframe(df_steps, use_container_width=True, hide_index=True)

            # --- PUBLICACIÓN Y EXPORTACIÓN ---
            st.markdown("---")
            st.markdown("### 🚀 Generar y Registrar Documentación")
            
            # Mensaje descriptivo dinámico según el tipo seleccionado
            if selected_type == "Flujo de Trabajo":
                desc_text = f"El flujo de trabajo se publicará formalmente bajo la versión <b>{input_version}</b> y fecha de aprobación <b>{fecha_proceso.strftime('%d/%m/%Y')}</b>."
            elif selected_type == "Norma":
                desc_text = f"El documento formal de la Norma se generará bajo la autoría de <b>{input_autor}</b>, versión <b>{input_version}</b> y fecha de aprobación <b>{fecha_proceso.strftime('%d/%m/%Y')}</b>."
            else:
                desc_text = f"El procedimiento formal y su flujograma se generarán bajo la autoría de <b>{input_autor}</b>, versión <b>{input_version}</b> y fecha de aprobación <b>{fecha_proceso.strftime('%d/%m/%Y')}</b>."
            
            st.markdown(f"<p style='color:#94A3B8;'>{desc_text}</p>", unsafe_allow_html=True)

            # Ajustar etiqueta del botón y ayuda emergente según el tipo de documento seleccionado
            if selected_type == "Flujo de Trabajo":
                btn_label = "💾 Registrar Flujo y Publicar en Drive"
                btn_help = "Genera y publica únicamente el diagrama del flujo y su registro en el inventario de la nube."
            elif selected_type == "Norma":
                btn_label = "💾 Generar Norma y Publicar en Drive"
                btn_help = "Genera y publica únicamente el documento formal de la Norma (.docx) en Drive y su registro en la nube."
            else:
                btn_label = "💾 Generar Procedimiento y Publicar en Drive"
                btn_help = "Genera el documento formal (.docx) y el flujograma de carriles, guardando y vinculando ambos en Drive y el inventario."
            
            btn_publish = st.button(btn_label, help=btn_help)
            
            if btn_publish:
                with st.spinner("Procesando y registrando documentación en la nube..."):
                    try:
                        fecha_str = fecha_proceso.strftime("%d/%m/%Y")
                        
                        # Usar el título editado manualmente si está disponible
                        title_to_use = st.session_state.get('document_title', doc.titulo)
                        if not title_to_use.strip():
                            title_to_use = doc.titulo
                            
                        child_proc = title_to_use
                        clean_title = "".join([c if c.isalnum() or c in [' ', '-', '_'] else '' for c in title_to_use]).strip()
                        
                        # Inicializar variables de estado
                        local_filepath = None
                        drawio_filepath = None
                        drive_link = None
                        flow_link = None
                        uploaded_doc_to_drive = False
                        uploaded_flow_to_drive = False
                        
                        # --- 1. GENERACIÓN DE WORD (.docx) ---
                        if selected_type != "Flujo de Trabajo":
                            data_dict = doc.model_dump()
                            data_dict["titulo"] = title_to_use # Sobrescribir con el título editado
                            data_dict["autor"] = input_autor
                            data_dict["version"] = input_version
                            data_dict["fecha"] = fecha_str
                            data_dict["codigo"] = code
                            
                            filename = f"{code} {clean_title}.docx"
                            
                            local_filepath = DocxGenerator.generate_document(
                                doc_type=selected_type,
                                data_dict=data_dict,
                                output_filename=filename
                            )
                            st.info(f"Documento formal .docx creado localmente en: {local_filepath}")
                            
                            # Subir a Google Drive
                            if os.path.exists(config.TOKEN_FILE):
                                try:
                                    drive_service = google_client.get_drive_service()
                                    file_metadata = {
                                        'name': filename,
                                        'parents': [config.DRIVE_FOLDER_ID]
                                    }
                                    from googleapiclient.http import MediaFileUpload
                                    media = MediaFileUpload(local_filepath, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                                    
                                    drive_file = drive_service.files().create(
                                        body=file_metadata,
                                        media_body=media,
                                        fields='id, webViewLink'
                                    ).execute()
                                    
                                    drive_link = drive_file.get('webViewLink')
                                    uploaded_doc_to_drive = True
                                    st.success(f"¡Documento Word subido a Google Drive de forma exitosa!")
                                except Exception as e_drive:
                                    st.warning(f"No se pudo subir el archivo Word a Google Drive automáticamente: {e_drive}")

                        # --- 2. GENERACIÓN DE DIAGRAMAS (SVG / DRAW.IO) ---
                        if selected_type in ["Procedimiento", "Instrucción de Trabajo", "Flujo de Trabajo"]:
                            # Forzar generación del diagrama local por si no se visitó la pestaña
                            svg_filepath = SVGGenerator.generate_svg(doc, code, input_version, fecha_str)
                            drawio_filepath = DrawIOGenerator.generate_drawio(doc, code, input_version, fecha_str)
                            
                            flow_filename = f"{code} {clean_title}.drawio"
                            
                            # Subir diagrama a Google Drive (.drawio)
                            if os.path.exists(config.TOKEN_FILE):
                                try:
                                    drive_service = google_client.get_drive_service()
                                    file_metadata = {
                                        'name': flow_filename,
                                        'parents': [config.DRIVE_FOLDER_ID]
                                    }
                                    from googleapiclient.http import MediaFileUpload
                                    media = MediaFileUpload(drawio_filepath, mimetype='application/octet-stream')
                                    
                                    drive_file = drive_service.files().create(
                                        body=file_metadata,
                                        media_body=media,
                                        fields='id, webViewLink'
                                    ).execute()
                                    
                                    flow_link = drive_file.get('webViewLink')
                                    uploaded_flow_to_drive = True
                                    st.success(f"¡Diagrama de flujo subido a Google Drive de forma exitosa!")
                                except Exception as e_flow_drive:
                                    st.warning(f"No se pudo subir el flujograma a Google Drive automáticamente: {e_flow_drive}")

                        # --- 3. REGISTRAR EN GOOGLE SHEETS / LOCAL CSV ---
                        parent_proc = title_to_use.replace("Procedimiento de ", "").replace("Norma de ", "")
                        
                        # Obtener el código relacionado para vincular proceso con flujo
                        rel_code = st.session_state.get('related_code', "")
                        
                        sheet_updated = im.save_new_document(
                            area_name=selected_area,
                            parent_proc=parent_proc,
                            child_proc=child_proc,
                            flow_link=flow_link or "",
                            code=code,
                            doc_type_name=selected_type,
                            doc_link=drive_link,
                            related_code=rel_code
                        )

                        # --- 4. RESUMEN DE RESULTADOS ---
                        st.markdown("### 🎉 Proceso Finalizado con Éxito")
                        
                        col_res_1, col_res_2 = st.columns(2)
                        with col_res_1:
                            st.markdown("#### 📂 Archivos en Local")
                            if local_filepath:
                                st.markdown(f"• **Documento formal:** `{os.path.basename(local_filepath)}`")
                                with open(local_filepath, "rb") as file:
                                    st.download_button(
                                        label="⬇️ Descargar archivo .docx",
                                        data=file,
                                        file_name=os.path.basename(local_filepath),
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="dl_docx_final"
                                    )
                            if drawio_filepath:
                                st.markdown(f"• **Flujograma editable:** `{os.path.basename(drawio_filepath)}`")
                                with open(drawio_filepath, "rb") as file:
                                    st.download_button(
                                        label="⚡ Descargar editable .drawio",
                                        data=file,
                                        file_name=os.path.basename(drawio_filepath),
                                        mime="application/octet-stream",
                                        key="dl_drawio_final"
                                    )
                            st.markdown(f"• **Carpeta local:** `{config.OUTPUTS_DIR}`")
                                    
                        with col_res_2:
                            st.markdown("#### ☁️ Registro en la Nube")
                            if uploaded_doc_to_drive and drive_link:
                                st.markdown(f"• **Enlace en Google Drive:** [Abrir Documento formal en la Nube]({drive_link})")
                            elif selected_type != "Flujo de Trabajo":
                                st.markdown("• **Google Drive:** No se pudo subir el documento formal.")
                                
                            if uploaded_flow_to_drive and flow_link:
                                st.markdown(f"• **Enlace en Google Drive:** [Abrir Flujograma en la Nube]({flow_link})")
                            elif selected_type in ["Procedimiento", "Instrucción de Trabajo", "Flujo de Trabajo"]:
                                st.markdown("• **Google Drive:** No se pudo subir el flujograma.")
                                
                            if sheet_updated:
                                st.markdown(f"• **Registro de Inventario:** Vinculado en tiempo real a la pestaña *Orden Matriz de Documentos*.")
                            else:
                                st.markdown(f"• **Registro de Inventario:** Guardado en el CSV local de salidas (`salidas/nuevos_documentos.csv`) para importar después.")
                                
                    except Exception as e_proc:
                        st.error(f"Error durante el proceso de emisión y guardado: {e_proc}")
                        import traceback
                        st.code(traceback.format_exc())
        else:
            # Elegant premium placeholder when no doc is loaded
            st.markdown("""
            <div style="
                background: #FFFFFF;
                border-radius: 16px;
                padding: 40px;
                border: 1.5px dashed #CBD5E1;
                text-align: center;
                margin-top: 50px;
                box-shadow: 0 4px 15px rgba(0,0,0,0.01);
            ">
                <div style="font-size: 4rem; margin-bottom: 20px;">⚡</div>
                <h3 style="color: #0F172A; font-family: 'Space Grotesk', sans-serif; font-weight: 700; margin-bottom: 10px; border:none; padding:0;">
                    Centro de Documentación Corporativa
                </h3>
                <p style="color: #64748B; font-size: 1.05rem; line-height: 1.6; max-width: 500px; margin: 0 auto 25px auto;">
                    Conversa con <b>Enterprise SGC Procesos IA</b> en el chat de la izquierda para diseñar tu proceso. Sube minutas de reuniones, PDFs o escribe notas de trabajo.
                </p>
                <div style="
                    background: rgba(248, 80, 0, 0.04);
                    border: 1px solid rgba(248, 80, 0, 0.15);
                    border-radius: 8px;
                    padding: 12px;
                    font-size: 0.9rem;
                    color: #F85000;
                    display: inline-block;
                    font-weight: 600;
                ">
                    💡 Haz clic en "Generar Estructura y Flujograma BPM" para ver el resultado aquí
                </div>
            </div>
            """, unsafe_allow_html=True)

with tab_matrix_inventory:
    st.markdown("### 🗂️ Inventario General de Documentos y Flujos (Farmacia Enterprise SGC)")
    st.markdown("<p style='color:#64748B;'>Esta matriz en tiempo real consolida todas las normas, procedimientos y flujos registrados en el Odoo/Sheets de Farmacia Enterprise SGC:</p>", unsafe_allow_html=True)
    
    # Función para limpiar hipervínculos de Sheets `=HYPERLINK("url", "text")` a "url"
    def clean_formula_to_url(val):
        if pd.isna(val):
            return ""
        val_str = str(val).strip()
        match = re.search(r'HYPERLINK\("([^"]+)"\s*,\s*"([^"]+)"\)', val_str, re.IGNORECASE)
        if match:
            return match.group(1)
        if val_str.startswith("http"):
            return val_str
        return val_str

    if im.matriz_df is not None and not im.matriz_df.empty:
        # Métricas rápidas estilo dashboard premium
        total_docs = len(im.matriz_df)
        
        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            st.markdown(f"""
            <div style="background-color:#FFFFFF; border:1px solid #E2E8F0; border-radius:10px; padding:15px; text-align:center; box-shadow: 0 4px 6px rgba(0,0,0,0.02);">
                <div style="font-size:0.85rem; font-weight:700; color:#475569; text-transform:uppercase; letter-spacing:0.5px;">Total Documentos</div>
                <div style="font-size:2.2rem; font-weight:800; color:#F85000; margin-top:5px;">{total_docs}</div>
            </div>
            """, unsafe_allow_html=True)
        with col_m2:
            areas_count = im.matriz_df['Area'].nunique() if 'Area' in im.matriz_df.columns else 0
            st.markdown(f"""
            <div style="background-color:#FFFFFF; border:1px solid #E2E8F0; border-radius:10px; padding:15px; text-align:center; box-shadow: 0 4px 6px rgba(0,0,0,0.02);">
                <div style="font-size:0.85rem; font-weight:700; color:#475569; text-transform:uppercase; letter-spacing:0.5px;">Áreas Operativas</div>
                <div style="font-size:2.2rem; font-weight:800; color:#475569; margin-top:5px;">{areas_count}</div>
            </div>
            """, unsafe_allow_html=True)
        with col_m3:
            # Contar procedimientos válidos
            proc_count = 0
            if 'Procedimiento' in im.matriz_df.columns:
                proc_count = im.matriz_df['Procedimiento'].apply(lambda x: str(x).strip() != '' and pd.notna(x) and str(x) != 'nan' and str(x) != '').sum()
            st.markdown(f"""
            <div style="background-color:#FFFFFF; border:1px solid #E2E8F0; border-radius:10px; padding:15px; text-align:center; box-shadow: 0 4px 6px rgba(0,0,0,0.02);">
                <div style="font-size:0.85rem; font-weight:700; color:#475569; text-transform:uppercase; letter-spacing:0.5px;">Procedimientos Activos</div>
                <div style="font-size:2.2rem; font-weight:800; color:#22C55E; margin-top:5px;">{proc_count}</div>
            </div>
            """, unsafe_allow_html=True)
            
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Buscador y filtros dinámicos compartidos
        search_query = st.text_input("🔍 Buscar por Palabra Clave (Título, Código o Área) en las matrices:", key="inventory_search_bar")
        
        # Sub-pestañas para organizar las diferentes matrices
        sub_tab1, sub_tab2 = st.tabs([
            "🟠 Matriz Final de Procesos y Flujos",
            "📄 Matriz Final de Normas"
        ])
        
        # 1. Pestaña Matriz Final de Procesos y Flujos
        with sub_tab1:
            st.markdown("#### Matriz Final de Procesos y Flujos BPM")
            if im.proc_flujos_df is not None and not im.proc_flujos_df.empty:
                filtered_pf = im.proc_flujos_df.copy()
                if search_query.strip() != "":
                    q = search_query.lower()
                    mask = filtered_pf.apply(lambda row: row.astype(str).str.lower().str.contains(q).any(), axis=1)
                    filtered_pf = filtered_pf[mask]
                
                # Limpiar celdas con fórmulas =HYPERLINK para pasarlas como URLs puras a Streamlit
                filtered_pf['Link flujo'] = filtered_pf['Link flujo'].apply(clean_formula_to_url)
                filtered_pf['Link proceso'] = filtered_pf['Link proceso'].apply(clean_formula_to_url)
                
                # Configurar columnas como hipervínculos clickeables en Streamlit
                st.dataframe(
                    filtered_pf,
                    column_config={
                        "Link flujo": st.column_config.LinkColumn("Link del Flujo (Drive)", display_text="Ver Flujo 🔗"),
                        "Link proceso": st.column_config.LinkColumn("Link del Proceso (Drive)", display_text="Ver Proceso 📄")
                    },
                    use_container_width=True,
                    hide_index=True
                )
            else:
                st.info("No se encontraron registros en la Matriz Final de Procesos y Flujos.")

        # 2. Pestaña Matriz Final de Normas
        with sub_tab2:
            st.markdown("#### Matriz Final de Normas y Políticas Corporativas")
            if im.normas_df is not None and not im.normas_df.empty:
                filtered_n = im.normas_df.copy()
                if search_query.strip() != "":
                    q = search_query.lower()
                    mask = filtered_n.apply(lambda row: row.astype(str).str.lower().str.contains(q).any(), axis=1)
                    filtered_n = filtered_n[mask]
                
                # Limpiar celdas con fórmulas =HYPERLINK
                filtered_n['Link documento'] = filtered_n['Link documento'].apply(clean_formula_to_url)
                
                # Configurar columna como hipervínculo clickeable
                st.dataframe(
                    filtered_n,
                    column_config={
                        "Link documento": st.column_config.LinkColumn("Link del Documento (Drive)", display_text="Ver Norma 📄")
                    },
                    use_container_width=True,
                    hide_index=True
                )
            else:
                st.info("No se encontraron registros en la Matriz Final de Normas.")

        st.info("💡 Consejo: Puedes filtrar y ordenar cualquier columna haciendo clic en sus cabeceras. Las actualizaciones del spreadsheet oficial se cargan al iniciar o al refrescar la página.")
    else:
        st.info("No se encontraron registros de matriz cargados en el inventario.")
